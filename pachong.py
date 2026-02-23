#https://www.69shuba.com/txt/81790/
#39061006

from bs4 import BeautifulSoup   #网页解析，数据获取
from docx import Document
import os
import re                       #正则表达式，文字匹配
import requests
from requests.adapters import HTTPAdapter
from http.cookiejar import CookieJar
import http.cookiejar as cookielib

findTitle = re.compile(r'<h1>(.*?)</h1>',re.S)
findEnd = re.compile(r'<p>(.*?)</p>',re.S)
findLink = re.compile(r'<a href="(.*?)">下一章</a>')


def main():
    # baseurl = input("请输入目录列表的那个网址(例如https://www.dusuu.com/ml/1033/)\n")
    # # https://www.dusuu.com/ml/10942/
    
    # number = input("请打开该小说第一章，输入一下网址新增加的那个数字(例如2688178)\n")
    # #2622647
    # savepath = "./"+input("请为你的docx文档命名:")+".docx"
    # url = baseurl+str(number)

    
    # test
    savepath = "./fuhuoquanrenlei.docx"
    url = 'https://twkan.com/txt/73947/45140676'

    Doc = Document()
    link = ""
    req=requests.Session()
    req.mount('https://',HTTPAdapter(max_retries=3))
    headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/145.0.0.0 Safari/537.36'
    }
    cookies = CookieJar()  # 或者从浏览器复制cookies到此变量中
    
    print("已完成爬取的章节:", url)
    while(link!="https://www.69shuba.com/book/81790.htm"):
        try:
            req = requests.get(url = url, headers = headers, cookies=cookies, timeout=5)
            req.encoding = 'utf-8'
            print('text', req.text)
            soup = BeautifulSoup(req.text, 'html5lib')
            for item in soup.find_all('h1',attrs = {'id':'txtnav'}):
                item = str(item)
                title = re.findall(findTitle, item)[0]
                print(title)
                Doc.add_heading(title, level = 0)
            for item in soup.find_all('div',attrs = {'id':'txtcontent0'}):
                item = str(item)
                end = re.findall(findEnd, item)
                for i in range(len(end)):
                    if(i==(len(end)-1)):
                        end[i]="\n"
                    Doc.add_paragraph(end[i])
            Doc.save(savepath)
            #跳转到下一页
            for item in soup.find_all('div', attrs = {'class':'page1'}):
                item = str(item)
                link = re.findall(findLink, item)[0]
                print('link', link)
                url = str(link)
        except requests.exceptions.RequestException as e:
            print(e)
    os.system("pause")

if __name__ == "__main__":
    main()
