from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.common import exceptions
from selenium.webdriver.remote.webdriver import WebElement
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import re
import os
import random
from bs4 import BeautifulSoup   #网页解析，数据获取
from bs4.element import Tag,PageElement
from docx import Document
import requests
from enum import Enum

class State(Enum):
    Init = 0
    Wait = 1
    GetUrl = 2
    ParseHtml = 3
    OverAutoer = 4
    End = 5

findTitle = re.compile(r'<h1>(.*?)</h1>',re.S)
findEnd = re.compile(r'<div>(.*?)</div>',re.S)
findLink = re.compile(r'<a href="(.*?)">下一章</a>')

def main():
    #word
    Doc = Document()

    """创建基础驱动"""
    chrome_options = Options()
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_argument("--disable-infobars")
    chrome_options.add_argument("--disable-extensions")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    
    # 禁用自动化特征
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)


    # 初始化Chrome驱动
    driver = webdriver.Chrome(options=chrome_options)

    # test
    current_path = os.path.dirname(os.path.realpath(__file__))
    savepath = current_path + "/架构整洁之道中文版 - Clean Architecture.docx"
    print('save path', savepath)
    url = 'https://geekdaxue.co/read/Clean-Architecture-zh/docs-ch2.md'

    #保存网址的队列
    urls = []
    print('begin')
    print('current url:',driver.current_url)

    _state:State = State.Init
    while(_state != State.End):
        try:
            print(_state)
            if _state == State.Init:
                driver.get(url)
                _state = State.Wait
                    
            elif _state == State.Wait:
                
                element:WebElement = driver.find_element(By.XPATH, '/html/body/div[1]/div[1]/div[2]/div[2]/div[1]/ul')
                if element:
                    links = element.find_elements(By.TAG_NAME, "a")
                    for item in links:
                        url_value = item.get_attribute("href")
                        print("打开网址：", url_value)
                        if url_value:
                            urls.append(url_value)
                
                    _state = State.OverAutoer

                else:
                    print("等待广告")
            elif _state == State.OverAutoer:
                print("通过广告")
                for item in urls:
                    driver.get(item)
                    while True:
                        title = driver.find_element(By.XPATH, '//*[@id="article-title"]')
                        if title:
                            print("title:", title.text)
                            Doc.add_heading(title.text, level = 0)
                        element = driver.find_element(By.XPATH, '/html/body/div[1]/div[1]/div[1]/div/div[2]/article')
                        if element:
                            #print(element.text)
                            Doc.add_paragraph(element.text)
                            #保存word
                            Doc.save(savepath)
                            break
                                        
                _state = State.End
            elif _state == State.ParseHtml:                
                pass
            
        except requests.exceptions.RequestException as e:
            print(e)
        except exceptions.NoSuchElementException as e:
            print(e)
            _state = State.Wait
            continue
        except exceptions.ElementNotInteractableException as e:
            print("element is not interactable")
            _state =State.Wait


    driver.quit()  # 关闭浏览器


if __name__ == "__main__":
    main()