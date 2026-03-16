from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.common import exceptions
from selenium.webdriver.remote.webdriver import WebElement
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import re
import os
import random
from bs4 import BeautifulSoup   #网页解析，数据获取
from bs4.element import Tag,PageElement
from docx import Document
import requests
from enum import Enum

class State(Enum):
    Init = 0
    Wait = 1
    GetUrl = 2
    ParseHtml = 3
    OverAutoer = 4

findTitle = re.compile(r'<h1>(.*?)</h1>',re.S)
findEnd = re.compile(r'<div>(.*?)</div>',re.S)
findLink = re.compile(r'<a href="(.*?)">下一章</a>')

def main():
    #word
    Doc = Document()

    """创建基础驱动"""
    chrome_options = Options()
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_argument("--disable-infobars")
    chrome_options.add_argument("--disable-extensions")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    
    # 禁用自动化特征
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)


    # 初始化Chrome驱动
    driver = webdriver.Chrome(options=chrome_options)

    # test
    current_path = os.path.dirname(os.path.realpath(__file__))
    savepath = current_path + "/fuhuoquanrenlei.docx"
    print('save path', savepath)
    url = 'https://www.69shuba.com/txt/81790/39061006'

    link = ""
    print('begin')
    print('current url:',driver.current_url)

    
    _state:State = State.Init
    _check_button:WebElement = None
    while(link!="https://www.69shuba.com/book/81790.htm"):
        try:
            print(_state)
            if _state == State.Init:
                driver.get(url)
                #time.sleep(5)  # 等待页面加载，根据需要调整时间
                page_source = driver.page_source  # 获取页面源代码
                #print(page_source)
                #//*[@id='CZUq4']/div/label/input
                _state = State.Wait
                    
            elif _state == State.Wait:
                
                time.sleep(random.randint(6,9))
                _check_button = driver.find_element(By.XPATH, '//*[@id="CZUq4"]/div/label/input')
                if _check_button:
                    print(_check_button)
                    if not _check_button.is_selected():
                        #_check_button.click()
                        webdriver.ActionChains(driver).move_to_element(_check_button).click(_check_button).perform()
                    _state = State.OverAutoer
            elif _state == State.GetUrl:
                pass
            elif _state == State.OverAutoer:
                time.sleep(1)
                _state = State.ParseHtml
                pass
            elif _state == State.ParseHtml:
                soup = BeautifulSoup(page_source, 'html5lib')
                titleElems = soup.find_all('h1',attrs = {'class':'hide720'})
                #标题
                for elem in titleElems:
                    title = elem.text
                    print('title:',title)
                    Doc.add_heading(title, level = 0)
                        
                        
                #内容
                txtElems = soup.find_all('div',attrs = {'class':'txtnav'})
                for elem in txtElems:
                    for item in elem.contents:
                        if not item.name:
                            Doc.add_paragraph(item.text)
                        #print(content)
                #保存word
                Doc.save(savepath)
                #跳转到下一页
                for item in soup.find_all('div', attrs = {'class':'page1'}):
                    item = str(item)
                    link = re.findall(findLink, item)[0]
                    print('link', link)
                    url = str(link)
            
            
        except requests.exceptions.RequestException as e:
            print(e)
        except exceptions.NoSuchElementException as e:
            print(e)
            _state = State.Wait
            continue
        except exceptions.ElementNotInteractableException as e:
            print("element is not interactable")
            _state =State.Wait


    driver.quit()  # 关闭浏览器


if __name__ == "__main__":
    main()